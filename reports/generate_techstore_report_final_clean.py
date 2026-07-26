# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
OUT = BASE / "BaoCao_TechStore_Final_KhongNhapHang.docx"
SCREENS = BASE / "techstore_report_v3_screens"
MYSQL = BASE / "techstore_report_v3_mysql"

BLUE = "1F4D78"
INK = "111827"
MUTED = "374151"
HEADER_FILL = "E8F0FA"


def set_run(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, INK),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True


def heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        set_run(r, {1: 16, 2: 13, 3: 12}.get(level, 11), True, color=BLUE if level < 3 else INK)
    return p


def para(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run(r)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run(r)


def numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        set_run(r)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def write_cell(cell_obj, text, bold=False, center=False):
    cell_obj.text = ""
    p = cell_obj.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    set_run(r, 10, bold)
    cell_obj.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, headers, rows, widths):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        write_cell(tbl.rows[0].cells[i], header, True, True)
        shade(tbl.rows[0].cells[i], HEADER_FILL)
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            write_cell(cells[i], value, center=i == 0)
    for row in tbl.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
    doc.add_paragraph()


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run(r, 10, italic=True, color=MUTED)


def figure(doc, image_path, cap, width=6.55, page_break=True):
    path = Path(image_path)
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(path), width=Inches(width))
        caption(doc, cap)
    else:
        para(doc, f"[Thiếu ảnh: {path.name}]", WD_ALIGN_PARAGRAPH.CENTER)
    if page_break:
        doc.add_page_break()


def cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG\nKHOA CÔNG NGHỆ THÔNG TIN")
    set_run(r, 14, True)

    doc.add_paragraph("\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LẬP TRÌNH HƯỚNG ĐỐI TƯỢNG")
    set_run(r, 18, True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ĐỀ TÀI: HỆ THỐNG QUẢN LÝ CỬA HÀNG TECHSTORE")
    set_run(r, 16, True)

    doc.add_paragraph("\n\n")
    for line in [
        "Giảng viên hướng dẫn: ........................................................",
        "Thành viên thực hiện: ........................................................",
        "Mã sinh viên: ........................................................",
        "Lớp: ........................................................",
        "Hà Nội - 2026",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run(r, 12, line.startswith("Giảng viên"))
    doc.add_page_break()


def table_of_contents(doc):
    heading(doc, "MỤC LỤC", 1)
    for line in [
        "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI",
        "1.1. Lý do chọn đề tài",
        "1.2. Mục tiêu của đề tài",
        "1.3. Phạm vi của đề tài",
        "1.4. Đối tượng sử dụng",
        "1.5. Tổng quan chức năng của hệ thống",
        "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ SỬ DỤNG",
        "2.1. Cơ sở lý thuyết về lập trình hướng đối tượng",
        "2.2. Các đặc trưng của lập trình hướng đối tượng được áp dụng trong đề tài",
        "2.3. Kiến trúc MVC",
        "2.4. Công nghệ Java",
        "2.5. Java Swing",
        "2.6. Lưu trữ dữ liệu trong MySQL",
        "2.7. Kiểm tra dữ liệu đầu vào",
        "2.8. Xử lý ngoại lệ",
        "CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG",
        "CHƯƠNG 4. CÀI ĐẶT, GIAO DIỆN VÀ HƯỚNG DẪN SỬ DỤNG",
        "CHƯƠNG 5. DỮ LIỆU MYSQL THỰC TẾ",
        "CHƯƠNG 6. KẾT LUẬN",
    ]:
        para(doc, line, WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()


def chapter1(doc):
    heading(doc, "CHƯƠNG 1. TỔNG QUAN ĐỀ TÀI", 1)
    heading(doc, "1.1. Lý do chọn đề tài", 2)
    para(doc, "Trong hoạt động kinh doanh thiết bị công nghệ, cửa hàng cần quản lý nhiều nhóm sản phẩm như linh kiện máy tính, linh kiện điện thoại và phụ kiện. Mỗi sản phẩm có các thông tin quan trọng như mã sản phẩm, tên sản phẩm, hãng sản xuất, giá bán, số lượng tồn kho và loại sản phẩm. Khi số lượng mặt hàng tăng lên, việc theo dõi thủ công bằng sổ sách hoặc các file rời rạc dễ dẫn đến sai lệch dữ liệu, khó kiểm soát tồn kho và mất nhiều thời gian khi tra cứu.")
    para(doc, "Bên cạnh việc quản lý sản phẩm, cửa hàng còn phải xử lý nhiều nghiệp vụ phát sinh hằng ngày như bán hàng, lập hóa đơn và thống kê doanh thu. Các nghiệp vụ này có mối liên hệ trực tiếp với số lượng tồn kho. Nếu dữ liệu không được cập nhật đồng bộ, cửa hàng có thể gặp tình trạng bán vượt số lượng tồn, ghi nhận sai doanh thu hoặc khó xác định lịch sử bán hàng của từng sản phẩm.")
    para(doc, "Xuất phát từ nhu cầu đó, đề tài \"Hệ thống quản lý cửa hàng TechStore\" được thực hiện nhằm xây dựng một ứng dụng desktop hỗ trợ quản lý các thông tin cơ bản liên quan đến hoạt động của cửa hàng công nghệ. Chương trình được phát triển bằng ngôn ngữ Java, sử dụng Java Swing để xây dựng giao diện người dùng và MySQL để lưu trữ dữ liệu tập trung.")
    para(doc, "Bên cạnh mục tiêu xây dựng một phần mềm có các chức năng quản lý thực tế, đề tài còn được thực hiện nhằm vận dụng các kiến thức của môn Lập trình Hướng đối tượng bằng Java. Trong quá trình xây dựng chương trình, các nguyên lý quan trọng của lập trình hướng đối tượng như tính đóng gói, tính kế thừa, tính đa hình và tính trừu tượng được áp dụng vào việc thiết kế các lớp model và xử lý nghiệp vụ.")
    para(doc, "Ngoài ra, hệ thống được tổ chức theo kiến trúc MVC, kết hợp với Service Layer và Repository. Cách tổ chức này giúp phân chia các thành phần của chương trình theo từng nhiệm vụ riêng biệt: giao diện hiển thị và nhận thao tác người dùng, controller điều phối sự kiện, service xử lý nghiệp vụ, repository làm việc với cơ sở dữ liệu. Nhờ đó, mã nguồn có cấu trúc rõ ràng, dễ quản lý và thuận tiện cho việc phát triển, bảo trì hệ thống.")
    para(doc, "Vì những lý do trên, nhóm lựa chọn đề tài \"Hệ thống quản lý cửa hàng TechStore\" làm bài tập lớn môn Lập trình Hướng đối tượng. Đề tài vừa phù hợp với kiến thức Java Desktop đã học, vừa có tính ứng dụng thực tế trong bài toán quản lý sản phẩm, bán hàng và thống kê doanh thu.")

    heading(doc, "1.2. Mục tiêu của đề tài", 2)
    para(doc, "Đề tài được thực hiện với mục tiêu xây dựng một phần mềm quản lý cửa hàng có giao diện trực quan, đáp ứng các chức năng quản lý cơ bản và đồng thời thể hiện được việc áp dụng các kiến thức về lập trình hướng đối tượng.")
    para(doc, "Các mục tiêu chính của đề tài bao gồm:")
    bullets(doc, [
        "Xây dựng một ứng dụng quản lý cửa hàng trên nền tảng Java Desktop.",
        "Sử dụng Java Swing để thiết kế giao diện người dùng.",
        "Áp dụng các nguyên lý cơ bản của lập trình hướng đối tượng vào quá trình xây dựng hệ thống.",
        "Tổ chức chương trình theo kiến trúc MVC, kết hợp giữa tầng giao diện, tầng xử lý nghiệp vụ và tầng lưu trữ dữ liệu.",
        "Xây dựng chức năng đăng nhập và phân quyền người dùng.",
        "Hỗ trợ quản lý thông tin sản phẩm, bao gồm linh kiện máy tính, linh kiện điện thoại và phụ kiện.",
        "Hỗ trợ lập hóa đơn bán hàng, chọn sản phẩm, kiểm tra số lượng tồn và lưu chi tiết hóa đơn.",
        "Hỗ trợ thống kê doanh thu, số lượng bán và hóa đơn liên quan theo sản phẩm.",
        "Kiểm tra tính hợp lệ của dữ liệu đầu vào nhằm hạn chế dữ liệu không chính xác.",
        "Xử lý các trường hợp ngoại lệ trong quá trình chương trình hoạt động.",
        "Lưu trữ và quản lý dữ liệu tập trung thông qua hệ quản trị cơ sở dữ liệu MySQL.",
    ])
    para(doc, "Thông qua việc thực hiện đề tài, nhóm có cơ hội vận dụng kiến thức lý thuyết vào một bài toán quản lý thực tế, đồng thời nâng cao kỹ năng phân tích, thiết kế và xây dựng phần mềm bằng ngôn ngữ Java.")

    heading(doc, "1.3. Phạm vi của đề tài", 2)
    para(doc, "Phạm vi của đề tài tập trung vào việc xây dựng một phần mềm desktop phục vụ cho công tác quản lý cửa hàng và các thông tin liên quan trong môi trường kinh doanh thiết bị công nghệ.")
    para(doc, "Hệ thống bao gồm các nhóm chức năng chính sau:")
    heading(doc, "1.3.1. Quản lý tài khoản và phân quyền", 3)
    para(doc, "Hệ thống cung cấp chức năng đăng nhập nhằm xác thực người dùng trước khi truy cập vào chương trình. Bên cạnh đó, chương trình hỗ trợ cơ chế phân quyền người dùng với các vai trò được xây dựng trong hệ thống, đồng thời cho phép người dùng sử dụng các chức năng phù hợp với quyền được cấp.")
    heading(doc, "1.3.2. Quản lý sản phẩm", 3)
    para(doc, "Chức năng quản lý sản phẩm cho phép lưu trữ và quản lý thông tin hàng hóa trong hệ thống. Mỗi sản phẩm có các thông tin như mã sản phẩm, tên sản phẩm, hãng sản xuất, giá bán, số lượng tồn và loại sản phẩm. Dữ liệu sản phẩm được sử dụng để phục vụ bán hàng và thống kê doanh thu.")
    heading(doc, "1.3.3. Quản lý hóa đơn", 3)
    para(doc, "Chức năng quản lý hóa đơn hỗ trợ tạo hóa đơn bán hàng, lưu chi tiết sản phẩm đã bán và xem lại thông tin hóa đơn sau khi lưu. Khi lập hóa đơn, hệ thống kiểm tra số lượng tồn kho và tính tổng tiền dựa trên số lượng, đơn giá của từng sản phẩm.")
    heading(doc, "1.3.4. Quản lý thống kê", 3)
    para(doc, "Chức năng thống kê cho phép theo dõi doanh thu, số lượng bán và hóa đơn liên quan theo từng sản phẩm trong khoảng thời gian lựa chọn. Người dùng có thể lọc dữ liệu theo danh mục sản phẩm, khoảng thời gian và tiêu chí thống kê.")
    para(doc, "Nhìn chung, phạm vi của đề tài tập trung vào việc xây dựng một hệ thống quản lý cửa hàng ở quy mô ứng dụng desktop, sử dụng dữ liệu được lưu trữ tập trung trong MySQL. Đề tài chưa tập trung vào việc triển khai hệ thống trên môi trường web hoặc xây dựng hệ thống bán hàng trực tuyến.")

    heading(doc, "1.4. Đối tượng sử dụng", 2)
    para(doc, "Phần mềm được xây dựng với cơ chế đăng nhập và phân quyền người dùng. Theo thiết kế của chương trình, hệ thống có các vai trò người dùng như ADMIN và NHANVIEN.")
    para(doc, "Tùy thuộc vào quyền được cấp, người dùng có thể truy cập và thực hiện các chức năng quản lý phù hợp với vai trò của mình. Việc phân quyền giúp tăng tính kiểm soát trong quá trình sử dụng phần mềm và hạn chế thao tác dữ liệu không phù hợp.")
    bullets(doc, [
        "Quản trị viên (ADMIN): thực hiện các công việc quản lý hệ thống và dữ liệu theo quyền được cấp.",
        "Nhân viên bán hàng (NHANVIEN): sử dụng các chức năng quản lý liên quan đến công việc được phân quyền trong hệ thống.",
    ])

    heading(doc, "1.5. Tổng quan chức năng của hệ thống", 2)
    table(doc, ["STT", "Module", "Chức năng chính"], [
        [1, "Đăng nhập và phân quyền", "Xác thực tài khoản và kiểm soát quyền truy cập."],
        [2, "Quản lý sản phẩm", "Quản lý danh sách sản phẩm, thêm mới, sửa thông tin và xóa sản phẩm."],
        [3, "Quản lý hóa đơn", "Lập hóa đơn bán hàng, chọn sản phẩm, nhập số lượng và xem chi tiết hóa đơn."],
        [4, "Thống kê doanh thu", "Tổng hợp doanh thu, số lượng bán và hóa đơn liên quan theo sản phẩm."],
    ], [0.55, 1.75, 4.55])
    para(doc, "Các chức năng trên được phân chia thành các module tương ứng trong chương trình. Về mặt cấu trúc, các module được tổ chức thành các lớp Model để biểu diễn dữ liệu, các lớp Service để xử lý nghiệp vụ, các lớp Repository để truy cập dữ liệu và các lớp View để xây dựng giao diện người dùng.")
    para(doc, "Theo cấu trúc của dự án, hệ thống có các thành phần chính như SanPham, LinhKienPC, LinhKienDienThoai, PhuKien, HoaDon và ChiTietHoaDon thuộc tầng Model. Các nghiệp vụ tương ứng được xử lý thông qua các lớp Service như SanPhamService, HoaDonService và ThongKeService. Giao diện người dùng được xây dựng bằng Java Swing với các thành phần như LoginFrame, MainFrame, SanPhamFrame, HoaDonFrame và ThongKeFrame.")


def chapter2(doc):
    heading(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ SỬ DỤNG", 1)
    heading(doc, "2.1. Cơ sở lý thuyết về lập trình hướng đối tượng", 2)
    heading(doc, "2.1.1. Khái niệm lập trình hướng đối tượng", 3)
    para(doc, "Lập trình hướng đối tượng (Object-Oriented Programming - OOP) là phương pháp lập trình tổ chức chương trình xoay quanh các đối tượng và lớp. Trong đó, đối tượng đại diện cho các thực thể trong bài toán, còn lớp là khuôn mẫu dùng để mô tả các thuộc tính và hành vi chung của các đối tượng.")
    para(doc, "Việc áp dụng lập trình hướng đối tượng giúp chương trình có cấu trúc rõ ràng, dễ mở rộng và thuận tiện cho việc bảo trì. Các đối tượng trong hệ thống có thể tương tác với nhau thông qua các phương thức được định nghĩa trong lớp.")
    para(doc, "Trong đề tài Hệ thống quản lý cửa hàng TechStore, phương pháp lập trình hướng đối tượng được sử dụng để mô hình hóa các đối tượng trong bài toán quản lý như sản phẩm, hóa đơn và chi tiết hóa đơn. Các đối tượng này được biểu diễn thông qua các lớp tương ứng trong package model.")
    para(doc, "Dự án áp dụng bốn đặc trưng cơ bản của lập trình hướng đối tượng gồm:")
    bullets(doc, [
        "Tính đóng gói (Encapsulation).",
        "Tính kế thừa (Inheritance).",
        "Tính đa hình (Polymorphism).",
        "Tính trừu tượng (Abstraction).",
    ])
    para(doc, "Việc áp dụng các đặc trưng trên giúp chương trình thể hiện rõ bản chất của lập trình hướng đối tượng và tạo ra cấu trúc mã nguồn có tính tổ chức cao.")

    heading(doc, "2.2. Các đặc trưng của lập trình hướng đối tượng được áp dụng trong đề tài", 2)
    heading(doc, "2.2.1. Tính đóng gói (Encapsulation)", 3)
    para(doc, "Tính đóng gói là cơ chế cho phép tập hợp dữ liệu và các phương thức xử lý dữ liệu vào bên trong một lớp, đồng thời hạn chế việc truy cập trực tiếp từ bên ngoài. Trong Java, tính đóng gói thường được thực hiện bằng cách khai báo các thuộc tính với phạm vi truy cập private và cung cấp các phương thức getter, setter để truy cập hoặc thay đổi dữ liệu.")
    para(doc, "Trong đề tài, các thuộc tính của các lớp Model như SanPham, LinhKienPC, LinhKienDienThoai, PhuKien, HoaDon và ChiTietHoaDon được đóng gói bên trong các lớp tương ứng. Các thuộc tính được quản lý thông qua các phương thức getter và setter.")
    para(doc, "Bên cạnh đó, chương trình còn sử dụng các lớp service để kiểm tra tính hợp lệ của dữ liệu đầu vào. Một số dữ liệu được kiểm tra bao gồm mã sản phẩm, tên sản phẩm, hãng sản xuất, giá bán, số lượng, mã hóa đơn và tên khách hàng. Giá bán và số lượng được kiểm tra ở phạm vi phù hợp với nghiệp vụ bán hàng.")
    para(doc, "Việc áp dụng tính đóng gói giúp hạn chế việc thay đổi dữ liệu không hợp lệ từ bên ngoài đối tượng, đồng thời nâng cao tính an toàn và khả năng kiểm soát dữ liệu của hệ thống.")

    heading(doc, "2.2.2. Tính kế thừa (Inheritance)", 3)
    para(doc, "Tính kế thừa cho phép một lớp mới kế thừa các thuộc tính và phương thức từ một lớp đã tồn tại. Lớp kế thừa có thể sử dụng lại các thành phần của lớp cha và bổ sung thêm các thuộc tính hoặc phương thức riêng phù hợp với yêu cầu của bài toán.")
    para(doc, "Trong hệ thống, lớp SanPham được xây dựng dưới dạng lớp trừu tượng và chứa các thông tin sản phẩm chung như:")
    bullets(doc, ["Mã sản phẩm.", "Tên sản phẩm.", "Hãng sản xuất.", "Giá bán.", "Số lượng tồn.", "Loại sản phẩm."])
    para(doc, "Từ lớp SanPham, các lớp LinhKienPC, LinhKienDienThoai và PhuKien được xây dựng bằng cơ chế kế thừa.")
    para(doc, "Lớp LinhKienPC kế thừa các thông tin chung từ SanPham và đại diện cho nhóm linh kiện máy tính. Lớp LinhKienDienThoai kế thừa các thông tin chung và đại diện cho nhóm linh kiện điện thoại. Lớp PhuKien kế thừa các thông tin chung và mô tả nhóm sản phẩm phụ kiện.")
    para(doc, "Việc sử dụng kế thừa giúp hạn chế việc lặp lại mã nguồn, đồng thời thể hiện mối quan hệ giữa các đối tượng trong thực tế. Trong trường hợp này, các nhóm sản phẩm đều có thông tin cơ bản giống nhau nên có thể kế thừa từ lớp SanPham.")

    heading(doc, "2.2.3. Tính đa hình (Polymorphism)", 3)
    para(doc, "Tính đa hình cho phép cùng một phương thức hoặc giao diện có thể được triển khai theo những cách khác nhau tùy thuộc vào đối tượng cụ thể.")
    para(doc, "Trong đề tài, tính đa hình được thể hiện thông qua việc các lớp con của SanPham có thể được xử lý dưới dạng kiểu SanPham trong service, controller và repository. Nhờ vậy, chương trình có thể quản lý nhiều loại sản phẩm theo một cấu trúc chung nhưng vẫn giữ được đặc điểm riêng của từng loại sản phẩm.")
    para(doc, "Ngoài ra, các phương thức xử lý trong từng service như thêm sản phẩm, tạo hóa đơn và thống kê doanh thu cũng thể hiện cách tổ chức nghiệp vụ theo các luồng xử lý chung. Controller chỉ cần gọi service tương ứng, còn chi tiết xử lý được đóng gói bên trong từng lớp nghiệp vụ.")

    heading(doc, "2.2.4. Tính trừu tượng (Abstraction)", 3)
    para(doc, "Tính trừu tượng cho phép tập trung vào những đặc điểm và hành vi quan trọng của đối tượng, đồng thời ẩn đi các chi tiết triển khai không cần thiết đối với bên ngoài.")
    para(doc, "Trong đề tài, tính trừu tượng được thể hiện rõ nhất thông qua lớp trừu tượng SanPham. Lớp này định nghĩa những thông tin và hành vi chung cho mọi loại sản phẩm trong hệ thống.")
    para(doc, "Các lớp LinhKienPC, LinhKienDienThoai và PhuKien kế thừa từ SanPham và triển khai các đặc điểm cụ thể của riêng mình. Cách thiết kế này giúp hệ thống có một lớp cơ sở chung để quản lý các đối tượng sản phẩm.")
    para(doc, "Ngoài ra, hệ thống còn sử dụng service và repository như những tầng trừu tượng trong xử lý dữ liệu và nghiệp vụ. Việc tách riêng các tầng xử lý giúp phần giao diện không cần trực tiếp quan tâm đến cách dữ liệu được lưu trữ hoặc cách nghiệp vụ được triển khai bên trong.")
    para(doc, "Nhờ tính trừu tượng, các thành phần trong hệ thống được phân chia rõ ràng, giúp giảm sự phụ thuộc giữa các module và tạo điều kiện thuận lợi cho việc mở rộng chương trình.")

    heading(doc, "2.3. Kiến trúc MVC", 2)
    heading(doc, "2.3.1. Khái niệm kiến trúc MVC", 3)
    para(doc, "MVC (Model - View - Controller) là một mô hình kiến trúc thường được sử dụng để tổ chức cấu trúc của ứng dụng bằng cách phân chia chương trình thành các thành phần có nhiệm vụ khác nhau.")
    para(doc, "Mô hình MVC bao gồm ba thành phần chính:")
    bullets(doc, [
        "Model: Đại diện cho dữ liệu và các đối tượng nghiệp vụ của hệ thống.",
        "View: Cung cấp giao diện để người dùng tương tác với chương trình.",
        "Controller: Tiếp nhận thao tác từ người dùng và điều phối quá trình xử lý.",
    ])
    para(doc, "Việc phân chia chương trình theo mô hình MVC giúp tách biệt phần giao diện với phần xử lý nghiệp vụ và dữ liệu. Nhờ đó, việc phát triển và bảo trì hệ thống trở nên thuận tiện hơn.")
    heading(doc, "2.3.2. Áp dụng MVC trong đề tài", 3)
    para(doc, "Trong phần mềm quản lý cửa hàng TechStore, kiến trúc chương trình được tổ chức theo hướng MVC kết hợp với Service Layer và Repository.")
    para(doc, "Package model chứa các lớp đại diện cho các đối tượng dữ liệu của hệ thống, bao gồm SanPham, LinhKienPC, LinhKienDienThoai, PhuKien, HoaDon và ChiTietHoaDon.")
    para(doc, "Package view chứa các thành phần giao diện được xây dựng bằng Java Swing, bao gồm:")
    bullets(doc, [
        "LoginFrame.",
        "MainFrame.",
        "SanPhamFrame.",
        "HoaDonFrame.",
        "ThongKeFrame.",
    ])
    para(doc, "Các thành phần này chịu trách nhiệm hiển thị thông tin và cung cấp giao diện để người dùng tương tác với hệ thống.")
    para(doc, "Package controller tiếp nhận sự kiện từ giao diện và gọi đến các service tương ứng. Trong dự án, phần xử lý nghiệp vụ được tách thành các lớp controller như:")
    bullets(doc, [
        "LoginController.",
        "MainController.",
        "SanPhamController.",
        "HoaDonController.",
        "ThongKeController.",
    ])
    para(doc, "Các lớp controller đóng vai trò điều phối luồng xử lý giữa giao diện và tầng nghiệp vụ.")
    para(doc, "Package service đóng vai trò xử lý nghiệp vụ trung gian giữa controller và repository. Đây là nơi kiểm tra dữ liệu, tính toán các giá trị nghiệp vụ như tổng tiền hóa đơn, kiểm tra số lượng tồn kho trước khi bán, kiểm tra tính hợp lệ của dữ liệu trước khi ghi xuống cơ sở dữ liệu và tổng hợp số liệu thống kê.")
    para(doc, "Trong dự án, phần xử lý nghiệp vụ được tách thành các lớp service tương ứng với từng nhóm chức năng:")
    bullets(doc, [
        "AuthService.",
        "SanPhamService.",
        "HoaDonService.",
        "ThongKeService.",
    ])
    para(doc, "Package repository chứa các lớp truy cập dữ liệu, có nhiệm vụ làm việc trực tiếp với MySQL thông qua JDBC. Cấu trúc này giúp các tầng trong hệ thống tách biệt và dễ mở rộng.")
    para(doc, "Ngoài ra, hệ thống còn sử dụng package exception để xử lý các lỗi nghiệp vụ và lỗi hệ thống, ví dụ như TechStoreException, giúp chương trình thông báo lỗi rõ ràng và ổn định hơn trong quá trình sử dụng.")
    para(doc, "Cấu trúc tổng quát của hệ thống được tổ chức theo chiều dọc như sau:")

    heading(doc, "2.4. Công nghệ Java", 2)
    para(doc, "Java là ngôn ngữ lập trình hướng đối tượng được sử dụng để xây dựng toàn bộ ứng dụng. Java cung cấp các đặc điểm phù hợp với yêu cầu của đề tài như hỗ trợ lớp, đối tượng, kế thừa, đa hình, trừu tượng và đóng gói.")
    para(doc, "Trong dự án, Java được sử dụng để:")
    bullets(doc, [
        "Xây dựng các lớp Model.",
        "Xây dựng các lớp Controller.",
        "Xây dựng các lớp Service.",
        "Xây dựng các lớp Repository truy cập MySQL.",
        "Xử lý và kiểm tra dữ liệu.",
        "Xây dựng giao diện ứng dụng bằng Java Swing.",
        "Xử lý các nghiệp vụ quản lý sản phẩm, hóa đơn và thống kê.",
    ])
    para(doc, "Việc sử dụng Java giúp đề tài có thể vận dụng trực tiếp các kiến thức trọng tâm của môn Lập trình Hướng đối tượng vào một ứng dụng thực tế.")
    heading(doc, "2.5. Java Swing", 2)
    para(doc, "Java Swing là bộ thư viện giao diện đồ họa được sử dụng để xây dựng ứng dụng desktop bằng Java. Swing cung cấp nhiều thành phần giao diện như cửa sổ, nút bấm, ô nhập liệu, bảng dữ liệu và hộp thoại.")
    para(doc, "Trong đề tài, Java Swing được sử dụng để xây dựng giao diện của phần mềm quản lý cửa hàng TechStore. Các thành phần giao diện chính bao gồm:")
    bullets(doc, [
        "LoginFrame: giao diện đăng nhập.",
        "MainFrame: cửa sổ menu chính của chương trình.",
        "SanPhamFrame: giao diện quản lý sản phẩm.",
        "HoaDonFrame: giao diện quản lý hóa đơn.",
        "ThongKeFrame: giao diện thống kê doanh thu.",
    ])
    para(doc, "Ngoài ra, các hộp thoại như SanPhamDialog, HoaDonDialog, ChonSanPhamDialog, ChiTietHoaDonDialog và HoaDonLienQuanDialog được sử dụng để hỗ trợ thao tác nhập liệu và xem chi tiết dữ liệu.")

    heading(doc, "2.6. Lưu trữ dữ liệu trong MySQL", 2)
    para(doc, "Trong đề tài, dữ liệu được lưu trữ tập trung trong MySQL. Cơ sở dữ liệu techstore_db là nơi lưu toàn bộ thông tin quan trọng của hệ thống như tài khoản người dùng, sản phẩm, hóa đơn và chi tiết hóa đơn.")
    para(doc, "Việc sử dụng MySQL giúp hệ thống bảo đảm dữ liệu được quản lý tập trung, dễ truy vấn, dễ đồng bộ và phù hợp với các nghiệp vụ có liên hệ chặt chẽ như quản lý sản phẩm và bán hàng. Khi tạo hóa đơn, dữ liệu chi tiết được lưu vào các bảng liên quan và số lượng tồn kho được cập nhật theo nghiệp vụ bán hàng.")
    para(doc, "Các bảng dữ liệu chính trong hệ thống bao gồm:")
    bullets(doc, [
        "user: lưu tài khoản đăng nhập và vai trò người dùng.",
        "san_pham: lưu danh sách sản phẩm, giá bán, tồn kho và loại sản phẩm.",
        "hoa_don: lưu thông tin chung của hóa đơn bán hàng.",
        "chi_tiet_hoa_don: lưu danh sách sản phẩm trong từng hóa đơn.",
    ])
    para(doc, "So với việc lưu dữ liệu bằng file CSV thuần túy, MySQL có ưu thế hơn về tính toàn vẹn, khả năng truy vấn, khả năng mở rộng và tính đồng bộ khi nhiều nghiệp vụ cùng thao tác trên dữ liệu. Đây là lựa chọn phù hợp cho hệ thống quản lý cửa hàng có dữ liệu phát sinh liên tục.")
    para(doc, "Trong phạm vi phiên bản hiện tại, hệ thống tập trung vào quản lý sản phẩm, bán hàng và thống kê doanh thu. Toàn bộ dữ liệu phục vụ các chức năng này được đọc ghi thông qua repository bằng JDBC.")

    heading(doc, "2.7. Kiểm tra dữ liệu đầu vào", 2)
    para(doc, "Trong quá trình nhập dữ liệu, việc kiểm tra tính hợp lệ là một yếu tố quan trọng nhằm hạn chế dữ liệu sai hoặc không đầy đủ. Hệ thống xây dựng các lớp service để kiểm tra dữ liệu đầu vào trước khi lưu vào MySQL.")
    para(doc, "Một số thông tin được kiểm tra bao gồm mã sản phẩm, tên sản phẩm, hãng sản xuất, giá bán, số lượng, mã hóa đơn và tên khách hàng.")
    para(doc, "Đối với dữ liệu giá bán và số lượng, hệ thống kiểm tra để bảo đảm giá trị hợp lệ và phù hợp với nghiệp vụ bán hàng.")
    heading(doc, "2.8. Xử lý ngoại lệ", 2)
    para(doc, "Ngoại lệ là những tình huống bất thường có thể xảy ra trong quá trình chương trình hoạt động. Nếu không được xử lý phù hợp, các lỗi này có thể khiến chương trình dừng đột ngột hoặc gây ảnh hưởng đến trải nghiệm người dùng.")
    para(doc, "Trong đề tài, hệ thống xử lý các trường hợp như dữ liệu đầu vào không hợp lệ, mã sản phẩm hoặc hóa đơn trùng, không đủ số lượng tồn kho hoặc thiếu kết nối MySQL bằng cách hiển thị thông báo rõ ràng cho người dùng và tránh làm gián đoạn chương trình.")
    para(doc, "Việc kiểm soát ngoại lệ giúp ứng dụng ổn định hơn trong quá trình sử dụng thực tế, đồng thời hỗ trợ người dùng nhận biết lỗi và thao tác lại đúng cách.")


def chapter3(doc):
    heading(doc, "CHƯƠNG 3. PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1)
    heading(doc, "3.1. Yêu cầu chức năng", 2)
    para(doc, "Hệ thống cần đáp ứng các nghiệp vụ quản lý cửa hàng cơ bản. Mỗi chức năng có giao diện riêng, dữ liệu được lấy từ MySQL và các thao tác quan trọng được kiểm tra trước khi lưu.")
    table(doc, ["Nhóm yêu cầu", "Chi tiết"], [
        ["Đăng nhập", "Kiểm tra tài khoản, mật khẩu trong bảng user trước khi mở menu chính."],
        ["Sản phẩm", "Hiển thị, thêm, sửa, xóa sản phẩm và phân loại sản phẩm."],
        ["Hóa đơn", "Tạo hóa đơn, chọn sản phẩm, kiểm tra số lượng tồn và lưu chi tiết hóa đơn."],
        ["Thống kê", "Tính doanh thu, số lượng bán và xem hóa đơn liên quan."],
    ], [1.75, 4.95])
    heading(doc, "3.2. Yêu cầu phi chức năng", 2)
    bullets(doc, [
        "Giao diện cần dễ nhìn, thao tác đơn giản và phản hồi lỗi rõ ràng.",
        "Dữ liệu đầu vào phải được kiểm tra để hạn chế lỗi khi lưu vào cơ sở dữ liệu.",
        "Các thao tác tạo hóa đơn cần bảo đảm cập nhật dữ liệu nhất quán.",
        "Mã nguồn cần chia package rõ ràng để dễ bảo trì và mở rộng.",
    ])
    heading(doc, "3.3. Thiết kế package", 2)
    table(doc, ["Package", "Vai trò"], [
        ["com.techstore.model", "Chứa các lớp biểu diễn dữ liệu nghiệp vụ."],
        ["com.techstore.view", "Chứa các màn hình và hộp thoại Swing."],
        ["com.techstore.controller", "Xử lý sự kiện từ giao diện."],
        ["com.techstore.service", "Kiểm tra dữ liệu và xử lý nghiệp vụ."],
        ["com.techstore.repository", "Truy cập MySQL thông qua JDBC."],
        ["com.techstore.util", "Chứa SanPhamFactory để tạo sản phẩm theo loại."],
        ["com.techstore.exception", "Chứa ngoại lệ tùy chỉnh TechStoreException."],
    ], [2.15, 4.55])
    heading(doc, "3.4. Thiết kế lớp model", 2)
    para(doc, "Lớp SanPham đóng vai trò lớp cơ sở trừu tượng cho các loại sản phẩm. Các lớp LinhKienPC, LinhKienDienThoai và PhuKien kế thừa từ SanPham. Hóa đơn được biểu diễn bằng cặp lớp thông tin chung và chi tiết.")
    table(doc, ["Lớp", "Vai trò"], [
        ["SanPham", "Lớp trừu tượng chứa thông tin chung của sản phẩm."],
        ["LinhKienPC", "Đại diện sản phẩm thuộc nhóm linh kiện máy tính."],
        ["LinhKienDienThoai", "Đại diện sản phẩm thuộc nhóm linh kiện điện thoại."],
        ["PhuKien", "Đại diện sản phẩm phụ kiện."],
        ["HoaDon", "Thông tin hóa đơn bán hàng và danh sách chi tiết."],
        ["ChiTietHoaDon", "Một dòng sản phẩm trong hóa đơn."],
    ], [2.1, 4.6])
    heading(doc, "3.5. Thiết kế cơ sở dữ liệu", 2)
    para(doc, "Cơ sở dữ liệu techstore_db gồm các bảng chính: user, san_pham, hoa_don và chi_tiet_hoa_don. Bảng chi tiết hóa đơn liên kết với bảng hóa đơn thông qua mã hóa đơn, đồng thời liên kết với sản phẩm qua mã sản phẩm.")
    table(doc, ["Bảng", "Nội dung"], [
        ["user", "Tài khoản đăng nhập và vai trò người dùng."],
        ["san_pham", "Danh sách sản phẩm, giá bán, tồn kho và loại sản phẩm."],
        ["hoa_don", "Thông tin chung của hóa đơn bán hàng."],
        ["chi_tiet_hoa_don", "Danh sách sản phẩm trong từng hóa đơn."],
    ], [2.05, 4.65])
    heading(doc, "3.6. Luồng xử lý tiêu biểu", 2)
    numbers(doc, [
        "Người dùng đăng nhập bằng tài khoản có trong bảng user.",
        "Controller nhận sự kiện từ giao diện và gọi service tương ứng.",
        "Service kiểm tra dữ liệu đầu vào, phát hiện dữ liệu thiếu hoặc không hợp lệ.",
        "Repository thực hiện truy vấn hoặc cập nhật dữ liệu MySQL.",
        "Kết quả được nạp lại lên bảng giao diện để người dùng quan sát.",
    ])
    doc.add_page_break()


def chapter4(doc):
    heading(doc, "CHƯƠNG 4. CÀI ĐẶT, GIAO DIỆN VÀ HƯỚNG DẪN SỬ DỤNG", 1)
    para(doc, "Chương này trình bày từng chức năng của hệ thống kèm ảnh chụp thật từ ứng dụng Java Swing đang chạy. Ảnh được chụp ở độ phân giải cao để khi đưa vào Word không bị vỡ hoặc quá mờ.")
    items = [
        ("4.1. Đăng nhập hệ thống", "Màn hình đăng nhập yêu cầu người dùng nhập tên đăng nhập và mật khẩu. Controller kiểm tra dữ liệu trong bảng user trước khi cho phép truy cập.", ["Mở chương trình Java.", "Nhập tên đăng nhập và mật khẩu.", "Nhấn Đăng nhập để vào menu chính."], "01_dang_nhap.png", "Hình 4.1. Màn hình đăng nhập"),
        ("4.2. Menu chính", "Sau khi đăng nhập thành công, người dùng được chuyển đến menu chính. Menu này là điểm điều hướng đến các phân hệ sản phẩm, hóa đơn và thống kê.", ["Chọn Quản lý sản phẩm để quản lý danh mục hàng hóa.", "Chọn Quản lý hóa đơn để lập và xem hóa đơn bán hàng.", "Chọn Thống kê doanh thu để xem kết quả kinh doanh."], "02_menu_chinh.png", "Hình 4.2. Menu chính"),
        ("4.3. Danh sách sản phẩm", "Màn hình quản lý sản phẩm hiển thị dữ liệu thật từ bảng san_pham trong MySQL.", ["Mở chức năng Quản lý sản phẩm.", "Quan sát danh sách sản phẩm đang có trong cơ sở dữ liệu.", "Chọn cột thao tác để sửa hoặc xóa sản phẩm."], "03_san_pham_danh_sach.png", "Hình 4.3. Danh sách sản phẩm lấy từ MySQL"),
        ("4.4. Thêm sản phẩm", "Form thêm sản phẩm mới cho phép nhập mã sản phẩm, tên, hãng sản xuất, giá bán, số lượng và loại sản phẩm.", ["Nhấn nút Thêm Sản Phẩm Mới.", "Nhập đầy đủ các trường bắt buộc.", "Nhấn Lưu để thêm sản phẩm."], "04_san_pham_them_moi.png", "Hình 4.4. Form thêm sản phẩm mới"),
        ("4.5. Sửa sản phẩm", "Chức năng sửa sản phẩm dùng lại form sản phẩm nhưng khóa mã sản phẩm vì đây là khóa định danh.", ["Chọn một sản phẩm trong bảng.", "Mở menu thao tác và chọn Sửa.", "Cập nhật tên, hãng, giá bán hoặc số lượng.", "Nhấn Lưu để cập nhật dữ liệu."], "05_san_pham_sua.png", "Hình 4.5. Form sửa thông tin sản phẩm"),
        ("4.6. Xóa sản phẩm", "Trước khi xóa, hệ thống hiển thị hộp thoại xác nhận để tránh thao tác nhầm.", ["Chọn sản phẩm cần xóa.", "Chọn thao tác Xóa.", "Đọc thông tin xác nhận.", "Chọn Yes để xóa hoặc No để hủy."], "06_san_pham_xoa.png", "Hình 4.6. Hộp thoại xác nhận xóa sản phẩm"),
        ("4.7. Menu thao tác sản phẩm", "Menu thao tác trên bảng giúp giao diện gọn hơn. Các lệnh sửa và xóa xuất hiện khi người dùng chọn dòng sản phẩm.", ["Chọn một dòng trong bảng sản phẩm.", "Nhấn vào cột thao tác.", "Chọn Sửa Sản Phẩm hoặc Xóa Sản Phẩm."], "07_san_pham_menu_thao_tac.png", "Hình 4.7. Menu thao tác trong màn hình sản phẩm"),
        ("4.8. Danh sách hóa đơn", "Màn hình hóa đơn hiển thị các hóa đơn đã lưu trong bảng hoa_don.", ["Mở chức năng Quản lý hóa đơn.", "Quan sát danh sách hóa đơn.", "Chọn thao tác để xem chi tiết hoặc xóa hóa đơn."], "08_hoa_don_danh_sach.png", "Hình 4.8. Danh sách hóa đơn"),
        ("4.9. Tạo hóa đơn bán hàng", "Form tạo hóa đơn cho phép nhập khách hàng, ngày tạo và danh sách sản phẩm trong giỏ hàng.", ["Nhấn Thêm Hóa Đơn Mới.", "Nhập tên khách hàng.", "Thêm sản phẩm vào giỏ hàng.", "Kiểm tra tổng tiền và nhấn Lưu."], "09_hoa_don_tao_moi.png", "Hình 4.9. Form tạo hóa đơn bán hàng"),
        ("4.10. Chọn sản phẩm cho hóa đơn", "Dialog chọn sản phẩm hỗ trợ tìm kiếm theo tên hoặc mã sản phẩm và lọc theo danh mục.", ["Nhấn Thêm Sản Phẩm trong form hóa đơn.", "Tìm kiếm hoặc lọc danh mục.", "Chọn một sản phẩm trong bảng.", "Nhập số lượng muốn mua."], "10_hoa_don_chon_san_pham.png", "Hình 4.10. Dialog tìm kiếm và chọn sản phẩm"),
        ("4.11. Xem chi tiết hóa đơn", "Chi tiết hóa đơn hiển thị thông tin khách hàng, ngày tạo, danh sách sản phẩm, số lượng, đơn giá, thành tiền và tổng tiền.", ["Chọn hóa đơn trong danh sách.", "Mở menu thao tác.", "Chọn Xem Chi Tiết Hóa Đơn."], "11_hoa_don_chi_tiet.png", "Hình 4.11. Dialog xem chi tiết hóa đơn"),
        ("4.12. Menu thao tác hóa đơn", "Menu thao tác hóa đơn cho phép xem chi tiết hoặc xóa hóa đơn.", ["Chọn dòng hóa đơn.", "Nhấn cột thao tác.", "Chọn Xem Chi Tiết hoặc Xóa Hóa Đơn."], "12_hoa_don_menu_thao_tac.png", "Hình 4.12. Menu thao tác hóa đơn"),
        ("4.13. Thống kê doanh thu", "Màn hình thống kê tổng hợp số lượng bán và doanh thu theo sản phẩm.", ["Mở Thống Kê Doanh Thu.", "Chọn danh mục hoặc khoảng thời gian.", "Chọn thống kê theo doanh thu hoặc số lượng bán.", "Nhấn Lọc Dữ Liệu để cập nhật kết quả."], "15_thong_ke_doanh_thu.png", "Hình 4.13. Màn hình thống kê doanh thu"),
        ("4.14. Hóa đơn liên quan trong thống kê", "Khi chọn một sản phẩm trong bảng thống kê, người dùng có thể xem danh sách hóa đơn đã mua sản phẩm đó.", ["Chọn sản phẩm trong bảng thống kê.", "Nhấn Xem Hóa Đơn Liên Quan.", "Quan sát các hóa đơn và tổng số lượng bán."], "16_thong_ke_hoa_don_lien_quan.png", "Hình 4.14. Danh sách hóa đơn liên quan"),
    ]
    for title, desc, steps, image, cap in items:
        heading(doc, title, 2)
        para(doc, desc)
        numbers(doc, steps)
        figure(doc, SCREENS / image, cap)


def chapter5(doc):
    heading(doc, "CHƯƠNG 5. DỮ LIỆU MYSQL THỰC TẾ", 1)
    para(doc, "Chương này trình bày ảnh dữ liệu lấy trực tiếp từ database MySQL local techstore_db thông qua kết nối JDBC. Toàn bộ ảnh phản ánh dữ liệu thật hiện có trong hệ thống tại thời điểm chụp.")
    para(doc, "Do máy không có mysql CLI trong PATH, dữ liệu được truy xuất bằng JDBC từ chính database mà ứng dụng sử dụng, sau đó xuất thành ảnh bảng để đưa vào báo cáo. Nội dung trong ảnh không phải dữ liệu tự tạo trong Word.")
    for name, cap in [
        ("01_mysql_danh_sach_bang.png", "Hình 5.1. Danh sách bảng trong database techstore_db"),
        ("02_mysql_user.png", "Hình 5.2. Dữ liệu bảng user"),
        ("03_mysql_san_pham.png", "Hình 5.3. Dữ liệu bảng san_pham"),
        ("04_mysql_hoa_don.png", "Hình 5.4. Dữ liệu bảng hoa_don"),
        ("05_mysql_chi_tiet_hoa_don.png", "Hình 5.5. Dữ liệu bảng chi_tiet_hoa_don"),
    ]:
        figure(doc, MYSQL / name, cap, width=6.75)


def chapter6(doc):
    heading(doc, "CHƯƠNG 6. KẾT LUẬN", 1)
    heading(doc, "6.1. Kết quả đạt được", 2)
    bullets(doc, [
        "Xây dựng được ứng dụng desktop quản lý cửa hàng TechStore bằng Java Swing.",
        "Kết nối và sử dụng MySQL để lưu trữ dữ liệu thật của hệ thống.",
        "Hoàn thiện các chức năng đăng nhập, quản lý sản phẩm, hóa đơn và thống kê.",
        "Tổ chức mã nguồn theo các tầng model, view, controller, service và repository.",
        "Áp dụng các đặc trưng lập trình hướng đối tượng trong thiết kế lớp và xử lý nghiệp vụ.",
    ])
    heading(doc, "6.2. Hạn chế", 2)
    bullets(doc, [
        "Giao diện Swing còn đơn giản, chưa có theme chuyên nghiệp.",
        "Chức năng phân quyền chưa chi tiết theo từng nghiệp vụ.",
        "Chưa có chức năng xuất báo cáo PDF/Excel trực tiếp từ ứng dụng.",
        "Chưa có kiểm thử tự động cho toàn bộ service và repository.",
    ])
    heading(doc, "6.3. Hướng phát triển", 2)
    bullets(doc, [
        "Bổ sung quản lý khách hàng, nhà cung cấp và bảo hành sản phẩm.",
        "Thêm biểu đồ thống kê trực quan và chức năng xuất báo cáo.",
        "Nâng cấp giao diện, cải thiện trải nghiệm người dùng.",
        "Bổ sung phân quyền người dùng và nhật ký thao tác.",
        "Triển khai kiểm thử tự động cho nghiệp vụ quan trọng.",
    ])
    heading(doc, "6.4. Tổng kết", 2)
    para(doc, "Đề tài Hệ thống quản lý cửa hàng TechStore đáp ứng mục tiêu của bài tập lớn môn Lập trình hướng đối tượng. Hệ thống có dữ liệu thật trong MySQL, giao diện desktop có thể thao tác được và mã nguồn được tổ chức theo kiến trúc rõ ràng. Đây là nền tảng phù hợp để tiếp tục mở rộng thành phần mềm quản lý cửa hàng hoàn chỉnh hơn.")


def main():
    doc = Document()
    configure(doc)
    cover(doc)
    table_of_contents(doc)
    chapter1(doc)
    doc.add_page_break()
    chapter2(doc)
    doc.add_page_break()
    chapter3(doc)
    chapter4(doc)
    chapter5(doc)
    chapter6(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
